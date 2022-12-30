import docx
from pathlib import Path

doc = docx.Document(Path('D:\_Python_Projects\Python-Docx\word_files', 'w_1.docx'))

print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(len(doc.paragraphs[0].runs))
print(doc.paragraphs[0].runs[0].text)

print('*' * 50)


def get_text(file_name):
    doc = docx.Document(file_name)
    full_text = [p.text for p in doc.paragraphs]
    return '\n'.join(full_text)


print(get_text(Path('D:\_Python_Projects\Python-Docx\word_files', 'w_1.docx')))
