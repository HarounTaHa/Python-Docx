import docx
from pathlib import Path
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

my_doc = docx.Document()

my_doc.add_paragraph('This is first paragraph of a MS Word file.')
my_doc.add_paragraph('This is Second paragraph of a MS Word file.')
# إضافة نص باللغة العربية
paragraph = my_doc.add_paragraph('التاريخ اليوم 14/1/2023')
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

third_paragraph = my_doc.add_paragraph('This is the third paragraph.')
third_paragraph.add_run('this is a section at the end of third paragraph. ')

# Add Heading
my_doc.add_heading('This is level 1', 0)
my_doc.add_heading('This is level 2', 1)
my_doc.add_heading('This is level 3', 2)
my_doc.add_heading('This is level 4', 3)

# Style
print(my_doc.paragraphs[0].style)
print(my_doc.paragraphs[5].text)
my_doc.paragraphs[0].style = my_doc.styles['Heading 1']

# Add Picture
my_doc.add_picture(str(Path.home() / Path('Desktop', 'logo.png')), width=docx.shared.Inches(5),
                   height=docx.shared.Inches(7))

my_doc.save(Path('D:\_Python_Projects\Python-Docx\word_files', 'write.docx'))
